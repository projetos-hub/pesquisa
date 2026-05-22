"""
Gera playbook DOCX com screenshots embutidas via python-docx.
v2: numeração manual (sem continuação), capa, quebras de página, note boxes.
Output: docs/playbook-visual.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT  = Path(__file__).parent.parent
SHOTS = ROOT / "docs" / "playbook-screenshots"
LOGO  = ROOT / "docs" / "logo-raiz.png"
OUT   = ROOT / "docs" / "playbook-visual.docx"

doc = Document()

# ── page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

NAVY   = RGBColor(0x0f, 0x34, 0x60)
DARK   = RGBColor(0x1a, 0x1a, 0x2e)
MID    = RGBColor(0x16, 0x21, 0x3e)
GRAY   = RGBColor(0x66, 0x66, 0x66)
AMBER  = RGBColor(0x92, 0x70, 0x00)
LGRAY  = RGBColor(0xaa, 0xaa, 0xaa)
RED    = RGBColor(0xe9, 0x45, 0x60)

# ── helpers ───────────────────────────────────────────────────────────────────
def page_break():
    doc.add_page_break()

def h1(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = DARK

def h2(text, new_page=True):
    if new_page:
        page_break()
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY

def h3(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = MID

def body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)

def bullet(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def numbered(items):
    """Numeração manual — evita continuação de lista entre seções."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent   = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        run = p.add_run(f"{i}. {item}")
        run.font.size = Pt(11)

def note(text, label="Importante"):
    """Caixa de nota com borda esquerda âmbar."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"{label}: ")
    run.bold  = True
    run.font.size = Pt(10)
    run.font.color.rgb = AMBER
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = AMBER
    # Borda esquerda via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), 'FFC107')
    pBdr.append(left)
    pPr.append(pBdr)
    # Fundo amarelo claro
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8E1')
    pPr.append(shd)
    doc.add_paragraph()

def tbl(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = NAVY
        # Fundo azul claro no header
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E8EEF7')
        tcPr.append(shd)
    # Data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        # Alternating rows
        if ri % 2 == 1:
            for ci in range(len(row)):
                tc = cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F8FAFF')
                tcPr.append(shd)
    doc.add_paragraph()

def screenshot(name, caption=""):
    f = SHOTS / f"{name}.png"
    if not f.exists():
        body(f"[screenshot {name} não encontrada]")
        return
    try:
        doc.add_picture(str(f), width=Inches(5.6))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(9)
            run.italic = True
            run.font.color.rgb = LGRAY
        doc.add_paragraph()
    except Exception as e:
        body(f"[erro ao inserir {name}: {e}]")

# ═══════════════════════════════════════════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

# Logo Raiz Educação
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(2.0))
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plataforma de Pesquisa CSAT")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Guia Operacional")
run.font.size = Pt(22)
run.font.color.rgb = NAVY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 30)
run.font.color.rgb = RED

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Raiz Educação  |  2026")
run.font.size = Pt(13)
run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Versão: 1.0  |  Plataforma: pesquisa-nu-sand.vercel.app")
run.font.size = Pt(10)
run.font.color.rgb = LGRAY

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUÇÃO
# ═══════════════════════════════════════════════════════════════════════════════
h2("1. Introdução", new_page=True)

h3("O que é o sistema")
body("A Plataforma de Pesquisa CSAT é um mini-app de pesquisa de satisfação integrado à Layers Education. Ela permite que a equipe da Raiz crie, configure e dispare pesquisas para alunos e responsáveis de escolas parceiras, coletando respostas de forma estruturada e exportando os dados para análise.")

h3("Personas")
body("Gestor Admin — membro da equipe Raiz com acesso ao painel administrativo. Cria e configura pesquisas, gerencia comunidades, dispara notificações e acompanha respostas.")
body("Respondente — aluno ou responsável de uma escola parceira. Acessa a pesquisa pelo portal da escola (Layers Education) ou por link direto e responde às perguntas.")

h3("Acesso ao sistema")
bullet([
    "Painel Admin: https://pesquisa-nu-sand.vercel.app/admin",
    "Login: e-mail e senha cadastrados no sistema",
    "Portal do respondente: integrado ao app Layers da escola via iframe",
])

# ═══════════════════════════════════════════════════════════════════════════════
# 2. PRIMEIROS PASSOS
# ═══════════════════════════════════════════════════════════════════════════════
h2("2. Primeiros Passos (Admin)")

h3("Login")
numbered([
    "Acesse https://pesquisa-nu-sand.vercel.app/admin/login",
    "Insira seu e-mail e senha",
    "Clique em Entrar",
])
body("Você será redirecionado automaticamente para a listagem de pesquisas.")

screenshot("01-login", "Tela de login — acesso ao painel administrativo")
screenshot("01b-login-preenchido", "Tela de login com credenciais preenchidas")

h3("Visão geral do painel")
body("O painel admin possui as seguintes seções no menu lateral:")
tbl(
    ["Seção", "Descrição"],
    [
        ["Pesquisas", "Lista e gerencia todas as pesquisas"],
        ["Comunidades", "Configura identidade visual global das escolas"],
        ["Disparos", "Cria e agenda notificações push/email para respondentes"],
        ["Exportar", "Faz download das respostas em XLSX"],
    ]
)

screenshot("02-painel-visao-geral", "Painel principal — listagem de pesquisas ativas e em rascunho")

# ═══════════════════════════════════════════════════════════════════════════════
# 3. GERENCIAMENTO DE PESQUISAS
# ═══════════════════════════════════════════════════════════════════════════════
h2("3. Gerenciamento de Pesquisas")

h3("Criar nova pesquisa")
numbered([
    "Acesse Pesquisas no menu",
    "Clique em + Nova pesquisa",
    "Preencha: Título, Slug (gerado automaticamente), Tipo, Datas de abertura/encerramento, Controle de acesso",
    "Clique em Criar pesquisa",
])

screenshot("03-criar-pesquisa-botao", "Botão '+ Nova pesquisa' em destaque no canto superior direito do painel")
screenshot("04-criar-pesquisa-form", "Formulário de criação de nova pesquisa")

h3("Painel da pesquisa")
body("Após criar, você acessa o painel individual da pesquisa, que centraliza todas as configurações: comunidades instaladas, amostra segmentada, disparos e respostas.")

screenshot("05-painel-pesquisa", "Painel da pesquisa — estatísticas, comunidades instaladas e ações rápidas")

h3("Status da pesquisa")
tbl(
    ["Status", "Significado"],
    [
        ["Rascunho", "Em construção, não visível para respondentes"],
        ["Ativa", "Aberta para respostas"],
        ["Pausada", "Temporariamente suspensa"],
        ["Encerrada", "Fechada, não aceita mais respostas"],
    ]
)
note("a pesquisa só fica acessível ao respondente quando o status é Ativa e a data de abertura já passou.")

screenshot("07-editar-pesquisa-status-datas", "Formulário de edição — status, datas de abertura e encerramento")

h3("Adicionar perguntas")
body("No painel da pesquisa, clique em Nova pergunta e escolha o tipo:")
tbl(
    ["Tipo", "Uso"],
    [
        ["NPS", "Nota de 0 a 10 com pergunta de recomendação"],
        ["Escala", "Avaliação de 1 a 5 para múltiplos aspectos"],
        ["Múltipla escolha (Radio)", "Seleção de uma opção entre várias"],
        ["Caixa de seleção (Checkbox)", "Seleção de múltiplas opções"],
        ["Texto aberto", "Campo livre para comentários"],
        ["Upload de arquivo", "Envio de documento, imagem, etc."],
    ]
)

screenshot("06-nova-pergunta-form", "Modal de nova pergunta — seleção de tipo e configuração dos campos")

# ═══════════════════════════════════════════════════════════════════════════════
# 4. COMUNIDADES E IDENTIDADE VISUAL
# ═══════════════════════════════════════════════════════════════════════════════
h2("4. Comunidades e Identidade Visual")

h3("O que é uma comunidade")
body("Uma comunidade representa uma escola parceira integrada à Layers Education. Cada escola tem um identificador único (community_id) que o Layers usa para identificar o contexto do usuário.")

h3("Instalar pesquisa em uma comunidade")
body("No painel da pesquisa, na seção Comunidades, insira o ID da escola no campo e clique em Instalar. Uma pesquisa só aparece para os respondentes das comunidades em que está instalada.")

screenshot("08-instalar-comunidade", "Seção Comunidades — campo de ID e botão Instalar para vincular escolas à pesquisa")

h3("Identidade visual")
body("A identidade visual (logo, cor primária, cor secundária) é configurada em dois níveis:")
bullet([
    "Nível global — em Identidade Visual no menu principal: vale para todas as pesquisas da escola (fallback)",
    "Nível por pesquisa — na aba Identidade Visual dentro de cada pesquisa: sobrescreve o tema global para aquela pesquisa",
])

screenshot("09-identidade-visual", "Identidade Visual por pesquisa — lista de comunidades com configuração de tema individual")
screenshot("16-comunidades-global", "Identidade Visual global — gerenciamento de temas de todas as escolas")
screenshot("17-editor-tema-global", "Editor de tema — cor primária, cor secundária, logo e preview em tempo real")

h3("Como configurar")
numbered([
    "Acesse a comunidade desejada e clique em Editar",
    "Faça upload do logo (PNG ou SVG recomendado)",
    "Defina as cores primária e secundária em hexadecimal (ex: #C8102E)",
    "Salve — o tema é aplicado imediatamente para os respondentes dessa escola",
])

# ═══════════════════════════════════════════════════════════════════════════════
# 5. DISPAROS
# ═══════════════════════════════════════════════════════════════════════════════
h2("5. Disparos de Notificação")

h3("Como funcionam")
body("O sistema envia notificações (push e/ou e-mail) para os usuários da escola no Layers. Os disparos são criados no painel e executados automaticamente pelo cron a cada 5 minutos.")

screenshot("10-disparos-visao-geral", "Painel de Disparos — formulário de novo disparo com destinatários, canais e mensagem")

h3("Régua de disparos (sequência automática)")
body("A régua é uma sequência de notificações programadas automaticamente a partir da data de abertura da pesquisa.")
numbered([
    "Acesse a pesquisa → aba Disparos",
    "Configure o conteúdo padrão (título e corpo da mensagem)",
    "Ative a Régua de disparos e adicione passos com offset em dias (ex: Passo 0 = dia da abertura, Passo 7 = 7 dias após)",
    "Escolha os canais: Push, E-mail ou ambos",
    "Clique em Criar régua",
])

screenshot("11-disparos-regua-form", "Formulário de disparo — seções de destinatários, canais e mensagem padrão")
screenshot("12-disparos-regua-passos", "Régua ativada — passos configurados: Convite inicial (Dia 0) e Lembrete (Dia 7)")

note("o horário dos disparos segue a hora de abertura da pesquisa. Configure 09:00 ou 17:00 para atingir o usuário no momento certo.", label="Dica")

h3("Disparo Rápido")
body("Para envios imediatos (testes ou comunicados urgentes), use o Disparo rápido no final da página de Disparos. Escolha a comunidade, informe os e-mails e clique em Enviar agora.")

screenshot("13-disparo-rapido", "Disparo Rápido expandido — comunidade, lista de e-mails e campos de mensagem")

h3("Amostra Segmentada")
body("Permite restringir o disparo a uma lista específica de usuários (ex: apenas responsáveis de uma turma).")
numbered([
    "Prepare uma planilha Excel com coluna email (e opcionalmente nome)",
    "Na aba Amostra da pesquisa, faça o upload",
    "O sistema resolve automaticamente os IDs Layers de cada e-mail",
    "Na aba Disparos, selecione Escopo: Amostra ao criar o disparo",
])

screenshot("14-amostra-segmentada", "Amostra Segmentada — upload de lista de e-mails e status de resolução dos IDs Layers")

h3("Placeholders disponíveis")
body("Use nos títulos e corpos das notificações e nas mensagens de boas-vindas/agradecimento:")
tbl(
    ["Placeholder", "O que exibe", "Fallback se não disponível"],
    [
        ["{{nome}}", "Primeiro nome do responsável", '"você"'],
        ["{{nomeAluno}}", "Nome completo do aluno", '"seu filho(a)"'],
        ["{{nomeEscola}}", "Nome da escola", '"a escola"'],
        ["{{serie}}", "Série/turma do aluno", '"a turma"'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. RESPOSTAS
# ═══════════════════════════════════════════════════════════════════════════════
h2("6. Acompanhar Respostas")

h3("Tabela de respostas")
body("Na aba Respostas de cada pesquisa: visualize todas as respostas em tabela paginada, filtre por nome, perfil (aluno/responsável), escola, série e onda.")

screenshot("15-respostas-tabela", "Tabela de respostas — listagem paginada com filtros por escola, perfil e onda")

h3("Exportar para XLSX")
numbered([
    "Acesse Exportar no menu principal",
    "Selecione a pesquisa",
    "Clique em ⬇ Baixar XLSX",
])
body("O arquivo gerado segue o padrão Metabase: uma linha por respondente, colunas de data/perfil/escola/série/onda + uma coluna por pergunta. Pronto para análise em BI.")

screenshot("18-exportar", "Tela Exportar — seleção de pesquisa e download do XLSX estruturado")

# ═══════════════════════════════════════════════════════════════════════════════
# 7. FLUXO DO RESPONDENTE
# ═══════════════════════════════════════════════════════════════════════════════
h2("7. Fluxo do Respondente")

h3("Como o respondente acessa")
body("Via LayersPortal (padrão): o app da escola exibe a pesquisa como iframe. O Layers passa automaticamente os dados do usuário (userId, communityId) para o sistema.")
body("Via link direto: o gestor pode compartilhar um link com parâmetros de comunidade para testes ou acesso alternativo.")

screenshot("19-respondente-portal", "Portal do respondente — tela de boas-vindas com personalização por nome e escola")

h3("Personalização automática")
body("Ao entrar, o sistema busca em tempo real: nome do responsável ou aluno, nome do filho (para responsáveis com vínculo ativo) e série/turma (via matrícula ativa).")

note("para {{serie}} aparecer, o aluno precisa ter uma matrícula ativa em uma turma na plataforma Layers — não basta estar em um grupo.")

h3("Percurso do respondente")
numbered([
    "Boas-vindas — mensagem de apresentação com nome personalizado",
    "Perguntas — percorre cada pergunta com barra de progresso",
    "Obrigatórias — não é possível avançar sem responder os campos obrigatórios",
    "Agradecimento — confirmação de envio com mensagem personalizada",
])

h3("Estados da pesquisa para o respondente")
tbl(
    ["Estado", "O que vê"],
    [
        ["Pesquisa não encontrada", "Tela de erro"],
        ["Ainda não aberta", '"Esta pesquisa ainda não está disponível"'],
        ["Encerrada", '"Esta pesquisa foi encerrada"'],
        ["Acesso negado", '"Você não tem acesso a esta pesquisa" (role incorreto)'],
        ["Aberta", "Fluxo normal de resposta"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. GLOSSÁRIO
# ═══════════════════════════════════════════════════════════════════════════════
h2("8. Glossário")

tbl(
    ["Termo", "Definição"],
    [
        ["Comunidade", "Escola parceira integrada à Layers Education, identificada por um community_id único"],
        ["Onda", 'Ciclo de uma pesquisa (ex: "1º semestre 2026") — campo informativo nas respostas'],
        ["Amostra", "Lista segmentada de usuários que podem responder uma pesquisa restrita"],
        ["Régua", "Sequência automática de notificações distribuídas ao longo do tempo"],
        ["Slug", "Identificador único da pesquisa na URL (ex: amostral-1-2026)"],
        ["Placeholder", "Variável no texto substituída por dado real do usuário (ex: {{nomeAluno}})"],
        ["Role", "Papel do usuário na plataforma Layers: guardian (responsável), student (aluno)"],
        ["CSAT", "Customer Satisfaction Score — tipo de pesquisa de satisfação"],
        ["NPS", "Net Promoter Score — pergunta de recomendação com nota de 0 a 10"],
        ["Offset", "Número de dias a partir da abertura para disparar uma notificação da régua"],
    ]
)

# ── rodapé ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph("Documento gerado automaticamente a partir do código-fonte da plataforma — Raiz Educação, 2026.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = LGRAY

# ── salvar ────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
size_kb = OUT.stat().st_size / 1024
print(f"OK: {OUT}")
print(f"   Tamanho: {size_kb:.0f} KB")
